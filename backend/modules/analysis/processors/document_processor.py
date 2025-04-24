"""
Procesador de documentos para el sistema de análisis
"""
from typing import Dict, Any, List, Optional, Union, TypeVar, Generic
import uuid
import logging
from sqlalchemy.orm import Session
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from datetime import datetime
from database.models.document import Document
from database.models.analysis import AnalysisPipeline
from core.llm_interface import LLMClientInterface
from docx import Document as DocxDocument
import pypdf
from io import BytesIO
import asyncio

logger = logging.getLogger(__name__)

# Definir un tipo genérico para las sesiones de base de datos
T = TypeVar('T', Session, AsyncSession)

class DocumentProcessor(Generic[T]):
    """Procesador para la extracción y procesamiento de documentos"""
    
    def __init__(self, llm_client: Optional[LLMClientInterface] = None):
        self.llm_client = llm_client
    
    def process_document(self, document: Document) -> Dict[str, Any]:
        """
        Procesar un documento y extraer su contenido
        
        Args:
            document: Documento a procesar
            
        Returns:
            Dict[str, Any]: Resultados del procesamiento
        """
        # Extraer texto del archivo
        text_content = self.extract_text_from_file(document)
        
        # Procesar contenido de texto (análisis básico)
        processing_results = self._process_text_content(text_content)
        
        # Retornar resultados
        return {
            "document_id": document.id,
            "text_content": text_content,
            "processing_results": processing_results,
            "processed_at": datetime.utcnow().isoformat()
        }
    
    def extract_text_from_file(self, document: Document) -> str:
        """
        Extraer texto de un archivo
        
        Args:
            document: Documento
            
        Returns:
            str: Contenido de texto
        """
        logger.info(f"Extrayendo texto del archivo: {document.file_path}")
        
        try:
            # Leer contenido del archivo como bytes
            with open(document.file_path, 'rb') as file:
                file_content = file.read()
            
            # Procesar según el tipo de archivo
            if document.file_path.lower().endswith(".docx"):
                return self._process_docx(file_content)
            elif document.file_path.lower().endswith(".pdf"):
                return self._process_pdf(file_content)
            else:
                # Para archivos de texto plano
                if document.file_path.lower().endswith((".txt", ".md", ".csv")):
                    with open(document.file_path, 'r', encoding='utf-8') as file:
                        return file.read()
                else:
                    raise ValueError(f"Tipo de archivo no soportado: {document.file_path}")
        except Exception as e:
            logger.error(f"Error al extraer texto del archivo {document.file_path}: {e}")
            raise
    
    def save_processing_results_sync(
        self, 
        db: Session, 
        pipeline_id: uuid.UUID, 
        document_id: uuid.UUID, 
        results: Dict[str, Any]
    ) -> Dict[str, Any]:
        """
        Guardar resultados del procesamiento en el pipeline de análisis (versión síncrona para Celery)
        
        Args:
            db: Sesión de base de datos síncrona
            pipeline_id: ID del pipeline de análisis
            document_id: ID del documento
            results: Resultados del procesamiento
            
        Returns:
            Dict[str, Any]: Resultados del procesamiento
        """
        # Obtener documento y pipeline usando API síncrona
        document = db.query(Document).filter(Document.id == document_id).first()
        pipeline = db.query(AnalysisPipeline).filter(AnalysisPipeline.id == pipeline_id).first()
        
        if not document:
            raise ValueError(f"Documento {document_id} no encontrado")
        if not pipeline:
            raise ValueError(f"Pipeline {pipeline_id} no encontrado")
        
        # Serializar resultados para JSON
        from utils.serialization import serialize_for_json
        serialized_results = serialize_for_json(results)
        
        # Actualizar metadatos del documento
        if document.process_metadata is None:
            document.process_metadata = {}
            
        document.process_metadata.update({
            "pipeline_processing": {
                "pipeline_id": str(pipeline_id),
                "processed_at": datetime.utcnow().isoformat(),
                "text_extracted": True
            }
        })
        
        # Actualizar resultados del pipeline
        processing_results = serialized_results.get("processing_results", {})
        if pipeline.processing_metadata is None:
            pipeline.processing_metadata = {}
            
        pipeline.processing_metadata.update({
            "document_processing": {
                "processed_at": datetime.utcnow().isoformat(),
                "text_extracted": True,
                "embeddings_generated": len(serialized_results.get("embeddings", [])) > 0,
                **processing_results
            }
        })
        
        # Guardar cambios con API síncrona
        db.commit()
        db.refresh(document)
        db.refresh(pipeline)
        
        return results
        
    async def save_processing_results_async(
        self, 
        db: AsyncSession, 
        pipeline_id: uuid.UUID, 
        document_id: uuid.UUID, 
        results: Dict[str, Any]
    ) -> Dict[str, Any]:
        """
        Guardar resultados del procesamiento en el pipeline de análisis (versión asíncrona para FastAPI)
        
        Args:
            db: Sesión de base de datos asíncrona
            pipeline_id: ID del pipeline de análisis
            document_id: ID del documento
            results: Resultados del procesamiento
            
        Returns:
            Dict[str, Any]: Resultados del procesamiento
        """
        # Obtener documento y pipeline usando API asíncrona
        document_result = await db.execute(select(Document).filter(Document.id == document_id))
        document = document_result.scalar_one_or_none()
        
        pipeline_result = await db.execute(select(AnalysisPipeline).filter(AnalysisPipeline.id == pipeline_id))
        pipeline = pipeline_result.scalar_one_or_none()
        
        if not document:
            raise ValueError(f"Documento {document_id} no encontrado")
        if not pipeline:
            raise ValueError(f"Pipeline {pipeline_id} no encontrado")
        
        # Serializar resultados para JSON
        from utils.serialization import serialize_for_json
        serialized_results = serialize_for_json(results)
        
        # Actualizar metadatos del documento
        if document.process_metadata is None:
            document.process_metadata = {}
            
        document.process_metadata.update({
            "pipeline_processing": {
                "pipeline_id": str(pipeline_id),
                "processed_at": datetime.utcnow().isoformat(),
                "text_extracted": True
            }
        })
        
        # Actualizar resultados del pipeline
        processing_results = serialized_results.get("processing_results", {})
        if pipeline.processing_metadata is None:
            pipeline.processing_metadata = {}
            
        pipeline.processing_metadata.update({
            "document_processing": {
                "processed_at": datetime.utcnow().isoformat(),
                "text_extracted": True,
                "embeddings_generated": len(serialized_results.get("embeddings", [])) > 0,
                **processing_results
            }
        })
        
        # Guardar cambios con API asíncrona
        await db.commit()
        await db.refresh(document)
        await db.refresh(pipeline)
        
        return results
        
    # Método de compatibilidad para código existente
    def save_processing_results(
        self, 
        db: Union[Session, AsyncSession], 
        pipeline_id: uuid.UUID, 
        document_id: uuid.UUID, 
        results: Dict[str, Any]
    ) -> Dict[str, Any]:
        """
        Guardar resultados del procesamiento en el pipeline de análisis
        (método de compatibilidad que detecta el tipo de sesión)
        
        Args:
            db: Sesión de base de datos (síncrona o asíncrona)
            pipeline_id: ID del pipeline de análisis
            document_id: ID del documento
            results: Resultados del procesamiento
            
        Returns:
            Dict[str, Any]: Resultados del procesamiento
        """
        if isinstance(db, AsyncSession):
            logger.warning("Usando save_processing_results con AsyncSession. Considere usar save_processing_results_async en su lugar.")
            # Ejecutar la versión asíncrona en un bucle de eventos
            return asyncio.run(self.save_processing_results_async(db, pipeline_id, document_id, results))
        else:
            return self.save_processing_results_sync(db, pipeline_id, document_id, results)
    
    def _process_text_content(self, text_content: str) -> Dict[str, Any]:
        """
        Procesar contenido de texto (análisis básico)
        
        Args:
            text_content: Contenido de texto
            
        Returns:
            Dict[str, Any]: Resultados del procesamiento
        """
        # Análisis básico del texto
        word_count = len(text_content.split())
        character_count = len(text_content)
        
        # Resultados del procesamiento
        return {
            "word_count": word_count,
            "character_count": character_count,
            "language": "es"  # Asumimos español por defecto
        }
    
    def _process_docx(self, file_content: bytes) -> str:
        """
        Procesar archivo DOCX
        
        Args:
            file_content: Contenido del archivo en bytes
            
        Returns:
            str: Texto extraído
        """
        doc = DocxDocument(BytesIO(file_content))
        text = ""
        
        # Extraer texto de párrafos
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n\n"
        
        # Extraer texto de tablas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
            text += "\n"
        
        return text
    
    def _process_pdf(self, file_content: bytes) -> str:
        """
        Procesar archivo PDF
        
        Args:
            file_content: Contenido del archivo en bytes
            
        Returns:
            str: Texto extraído
        """
        pdf_reader = pypdf.PdfReader(BytesIO(file_content))
        text = ""
        
        # Extraer texto de cada página
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text.strip():
                text += page_text + "\n\n"
        
        return text
